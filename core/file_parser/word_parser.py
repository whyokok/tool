"""
Microsoft Word file parser using python-docx.
"""
from pathlib import Path

from docx import Document

from .base import BaseParser, ParseResult


class WordParser(BaseParser):
    """Parser for Word documents (.docx)."""

    @property
    def file_type(self) -> str:
        return "word"

    def parse(self, file_path: Path) -> ParseResult:
        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            content = "\n".join(text_parts)
            return self._create_result(file_path, content.strip())

        except Exception as e:
            return self._create_result(file_path, "", str(e))
